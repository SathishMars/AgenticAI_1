import streamlit as st
import os
import tempfile
import pandas as pd
import pytesseract
from PIL import Image
from dotenv import load_dotenv
import pypandoc
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_groq import ChatGroq
from langchain.document_loaders import (
    PyPDFLoader, UnstructuredPowerPointLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.agents import Tool, initialize_agent, AgentType
from langchain.agents.agent import AgentExecutor
from langchain_experimental.agents.agent_toolkits.pandas.base import create_pandas_dataframe_agent

# Load API key
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

st.set_page_config(page_title="All-in-One Agentic Assistant", page_icon="🤖")
st.title("🧠 Multi-Document Agentic Assistant")
st.write("Upload PDF, CSV, Excel, Word (.doc/.docx), PPT, or image files and ask your questions.")

if not groq_api_key:
    st.error("🚫 GROQ_API_KEY not found in `.env`. Please add it before proceeding.")
    st.stop()

# Initialize Groq LLM
llm = ChatGroq(
    model_name="llama-3.3-70b-versatile",
    temperature=0.7,
    max_tokens=1024
)

# File uploader
uploaded_files = st.file_uploader(
    "Upload your files (PDF, Excel, CSV, Word, PPT, Images)",
    accept_multiple_files=True
)

tools = []

# Word loader that handles both .doc and .docx
def load_doc_or_docx_as_text(path):
    try:
        if path.endswith(".doc"):
            # Convert .doc to .docx using Pandoc
            converted_path = path + "x"
            pypandoc.convert_file(path, 'docx', outputfile=converted_path)
            path = converted_path

        # Load the .docx file content
        doc = DocxDocument(path)
        return "\n".join([para.text for para in doc.paragraphs])

    except Exception as e:
        raise RuntimeError(f"Failed to load Word file: {e}")

# Process uploaded files
for file in uploaded_files:
    filename = file.name.lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=filename) as tmp_file:
        tmp_file.write(file.read())
        tmp_path = tmp_file.name

    try:
        # === PDF ===
        if filename.endswith(".pdf"):
            st.info(f"📄 Processing PDF: {filename}")
            loader = PyPDFLoader(tmp_path)
            documents = loader.load()

        # === Word (.doc or .docx) ===
        elif filename.endswith((".doc", ".docx")):
            st.info(f"📝 Processing Word Document: {filename}")
            text = load_doc_or_docx_as_text(tmp_path)
            documents = [Document(page_content=text)]

        # === PowerPoint (.pptx) ===
        elif filename.endswith(".pptx"):
            st.info(f"📽️ Processing PowerPoint: {filename}")
            loader = UnstructuredPowerPointLoader(tmp_path)
            documents = loader.load()

        # === Image (JPG/PNG) ===
        elif filename.endswith((".jpg", ".jpeg", ".png")):
            st.info(f"🖼️ Processing Image: {filename}")
            img = Image.open(tmp_path)
            text = pytesseract.image_to_string(img)
            documents = [Document(page_content=text)]

        # === Excel/CSV ===
        elif filename.endswith((".csv", ".xls", ".xlsx")):
            st.info(f"📊 Processing Spreadsheet: {filename}")
            if filename.endswith(".csv"):
                df = pd.read_csv(tmp_path)
            else:
                df = pd.read_excel(tmp_path)

            agent = create_pandas_dataframe_agent(
                llm,
                df,
                verbose=False,
                allow_dangerous_code=True
            )
            df_tool = Tool(
                name=f"Spreadsheet - {filename}",
                func=agent.run,
                description=f"Use this to answer questions about spreadsheet {filename}."
            )
            tools.append(df_tool)
            continue  # Skip vectorizing spreadsheet

        else:
            st.warning(f"⚠️ Unsupported file type: {filename}")
            continue

        # === Vector-based QA for text documents ===
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = text_splitter.split_documents(documents)

        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vectorstore = FAISS.from_documents(docs, embeddings)
        retriever = vectorstore.as_retriever()

        tool = Tool(
            name=f"QA Tool - {filename}",
            func=RetrievalQA.from_chain_type(llm=llm, retriever=retriever),
            description=f"Use this to answer questions from {filename}."
        )
        tools.append(tool)

    except Exception as e:
        st.error(f"❌ Error processing {filename}: {e}")

# === Create Multi-tool Agent with parsing error handling ===
if tools:
    agent = initialize_agent(
        tools=tools,
        llm=llm,
        agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
        verbose=False,
        handle_parsing_errors=True  # ✅ This line fixes output parsing failures
    )

    st.markdown("### 💬 Ask a question about any of your uploaded files")
    query = st.text_input("Your question:")

    if query:
        with st.spinner("🔎 Thinking..."):
            try:
                result = agent.run(query)
                st.markdown("**🧠 Answer:**")
                st.write(result)
            except Exception as e:
                st.error(f"❌ Failed to answer: {e}")
else:
    st.warning("📂 Please upload at least one supported file.")
